# llm_service.py

import os
import pandas as pd
from io import StringIO
from docx import Document as DocxDocument
import asyncio
import datetime
from langchain.chains.history_aware_retriever import create_history_aware_retriever
from langchain.chains.retrieval import create_retrieval_chain
from langchain_core.messages import HumanMessage
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder, PromptTemplate
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyMuPDFLoader
from langchain.text_splitter import CharacterTextSplitter
from langchain.schema import Document
from langchain.chains.combine_documents import create_stuff_documents_chain
import tiktoken
from pymupdf.mupdf import ll_pdf_lookup_substitute_font_outparams

from db_service import DatabaseService
from settings import OPENAI_API_KEY, MODEL_NAME, CHAT_HISTORY_LEVEL, DOCS_IN_RETRIEVER
from helpers import current_timestamp


class LLMService:
    vector_store = None  # Class variable to store the vector store
    def __init__(self, model_name=MODEL_NAME):
        self.llm = ChatOpenAI(openai_api_key=OPENAI_API_KEY, model_name=model_name)
#        self.vector_store = None

    def load_excel_file(self, file_path):
        data = pd.read_excel(file_path)
        text_data = StringIO()
        data.to_string(buf=text_data)
        return text_data.getvalue()

    def load_word_file(self, file_path):
        doc = DocxDocument(file_path)
        return "\n".join([para.text for para in doc.paragraphs])


    def get_metadata(self, folder_path, db_service):
        from langchain.schema import HumanMessage
        metadata_list = []
        existing_file_paths = []

        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)

            # Skip directories
            if os.path.isdir(file_path):
                continue

            # Get the last modified time of the file from filesystem
            timestamp = os.path.getmtime(file_path)
            file_date_modified = datetime.datetime.fromtimestamp(timestamp)
            date_modify_str = file_date_modified.strftime('%Y-%m-%d %H:%M:%S')

            # Get dates from database
            db_date_modified, date_of_analysis = db_service.get_file_dates(file_path)

            if date_of_analysis:
                # Compare file's date_modified with date_of_analysis
                if file_date_modified <= date_of_analysis:
                    # File has not been modified since last analysis; skip processing
                    print(f"Skipping {filename}; no changes detected.")
                    continue
                else:
                    print(f"Re-analyzing {filename}; file has been modified.")
            else:
                print(f"Analyzing new file: {filename}")

            # Load content based on file type
            if filename.endswith(".pdf"):
                loader = PyMuPDFLoader(file_path)
                docs = loader.load()
                content = ' '.join([doc.page_content for doc in docs])

            elif filename.endswith(".docx"):
                content = self.load_word_file(file_path)

            elif filename.endswith(".xlsx"):
                content = self.load_excel_file(file_path)

            else:
                continue  # Skip unsupported file types

            # Limit content to first 2000 characters to avoid long prompts
            content_sample = content[:2000]

            # Generate AI description and document type
            prompt = (
                "Please analyze the following document content and provide the document type and a brief description in the following format:\n\n"
                "Document Type: [document type]\n"
                "Description: [description]\n\n"
                "Content:\n"
                f"{content_sample}"
            )

            response = self.llm.invoke([HumanMessage(content=prompt)]).content

            # Extract description and document type from response
            document_type = ''
            description = ''
            lines = response.strip().split('\n')
            for line in lines:
                if line.lower().startswith('document type:'):
                    document_type = line[len('document type:'):].strip()
                elif line.lower().startswith('description:'):
                    description = line[len('description:'):].strip()

            # Append metadata as a dictionary to the list
            metadata_list.append({
                'filename': filename,
                'path_file': file_path,
                'document_type': document_type,
                'date_modified': date_modify_str,
                'description': description
            })

            # After processing all files, mark files as deleted if they are not in the folder
            db_service.mark_files_as_deleted(existing_file_paths)

        return metadata_list


    def load_and_index_documents(self, folder_path):
        documents = []
        found_valid_file = False

        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)

            if filename.endswith(".pdf"):
                loader = PyMuPDFLoader(file_path)
                docs = loader.load()
                for doc in docs:
                    doc.metadata["source"] = filename
                    documents.append(doc)
                found_valid_file = True

            elif filename.endswith(".docx"):
                content = self.load_word_file(file_path)
                doc = Document(page_content=content, metadata={"source": filename})
                documents.append(doc)
                found_valid_file = True

            elif filename.endswith(".xlsx"):
                content = self.load_excel_file(file_path)
                doc = Document(page_content=content, metadata={"source": filename})
                documents.append(doc)
                found_valid_file = True

        if not found_valid_file:
            return "No valid files found in the folder. Please provide PDF, Word, or Excel files."

        text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        split_docs = text_splitter.split_documents(documents)

        embeddings = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)
        LLMService.vector_store = FAISS.from_documents(split_docs, embeddings)
        return "Documents successfully indexed."

    def generate_response(self, prompt, chat_history=None):

        if not LLMService.vector_store:
            return (
                "Please set the folder path using /folder and ensure documents are loaded.",
                None,
            )

        # Ensure chat_history is a list
        if chat_history is None:
            chat_history = []

        # Create the retriever with k=
        retriever = LLMService.vector_store.as_retriever(search_kwargs={'k': DOCS_IN_RETRIEVER})

        # Create the history-aware retriever
        retriever_prompt = ChatPromptTemplate.from_messages(
            [
                MessagesPlaceholder(variable_name="chat_history"),
                ("user", "{input}"),
                (
                    "user",
                    "Given the above conversation, generate a search query to look up in order to get information relevant to the conversation",
                ),
            ]
        )

        history_aware_retriever = create_history_aware_retriever(
            llm=self.llm, retriever=retriever, prompt=retriever_prompt
        )

        # Create the question-answering chain
        system_prompt = (
            "You are a project assistant from consultant side on design and construction projects."
            "Use the following pieces of retrieved context to answer "
            "the question. If you don't know the answer, say that you "
            " Do not include references to the source documents in your answer."
            f"don't know. If you need to use current date, today is {current_timestamp()}."
            "If Prompt include request to provide a link to documents in context, respond have to be: Please follow the link below:"
            " \n\n{context}"
        )

        prompt_template = ChatPromptTemplate.from_messages(
            [
                ("system", system_prompt),
                MessagesPlaceholder(variable_name="chat_history"),
                ("user", "{input}"),
            ]
        )

        question_answer_chain = create_stuff_documents_chain(self.llm, prompt_template)

        # Create the retrieval chain using create_retrieval_chain
        rag_chain = create_retrieval_chain(
            retriever=history_aware_retriever, combine_docs_chain=question_answer_chain
        )

        # Run the chain with the provided prompt and chat history
        result = rag_chain.invoke({"input": prompt, "chat_history": chat_history})

        answer = result.get("answer", "")
        sources = result.get("context", [])


        if not sources:
            return answer, None

        source_files = set(
            [doc.metadata["source"] for doc in sources if "source" in doc.metadata]
        )

        # Build the references
        references = {}
        for doc in sources:
            filename = doc.metadata.get('source', 'Unknown')
            page = doc.metadata.get('page', 'Unknown')
            if filename not in references:
                references[filename] = set()
            references[filename].add(page)

        return answer, source_files

    def count_tokens_in_context(self, folder_path):
        """Counts the total number of tokens in documents within a folder."""
        documents = []
        found_valid_file = False

        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)

            if filename.endswith(".pdf"):
                loader = PyMuPDFLoader(file_path)
                docs = loader.load()
                for doc in docs:
                    doc.metadata = {"source": filename}
                    documents.append(doc)
                found_valid_file = True

            elif filename.endswith(".docx"):
                content = self.load_word_file(file_path)
                doc = Document(page_content=content, metadata={"source": filename})
                documents.append(doc)
                found_valid_file = True

            elif filename.endswith(".xlsx"):
                content = self.load_excel_file(file_path)
                doc = Document(page_content=content, metadata={"source": filename})
                documents.append(doc)
                found_valid_file = True

        if not found_valid_file:
            return "No valid files found in the folder. Please provide PDF, Word, or Excel files."

        # Count tokens in the documents
        tokenizer = tiktoken.encoding_for_model(
            "gpt-4"
        )  # Tokenizer for the specific model

        total_tokens = 0
        for doc in documents:
            total_tokens += len(tokenizer.encode(doc.page_content))

        return total_tokens

